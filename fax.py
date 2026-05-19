import os
import uuid
import json
import sqlite3
import logging
import threading
import time
from datetime import datetime
from pathlib import Path
from io import BytesIO

import stripe
import telnyx
import requests
from flask import Flask, request, jsonify, redirect
from PIL import Image
import PyPDF2
from docx import Document
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

app = Flask(__name__)
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger(__name__)

# Admin panel
try:
    from fax_admin import register_fax_admin
    register_fax_admin(app, db_path="/home/ubuntu/harbor-fax.db",
                       file_dir="/tmp/harbor-fax-uploads")
except Exception as _e:
    log.exception("fax_admin failed to register: %s", _e)

STRIPE_SECRET           = os.environ["STRIPE_SECRET"]
STRIPE_WEBHOOK_SECRET   = os.environ.get("STRIPE_WEBHOOK_SECRET", "")
TELNYX_API_KEY          = os.environ["TELNYX_API_KEY"]
TELNYX_CONNECTION_ID    = os.environ["TELNYX_CONNECTION_ID"]
TELNYX_FROM_NUMBER      = os.environ["TELNYX_FROM_NUMBER"]
NTFY_TOPIC              = os.environ.get("NTFY_TOPIC", "harbor-brazer-monitor")
NTFY_AUTH               = os.environ.get("NTFY_AUTH", "")
RESEND_API_KEY          = os.environ.get("RESEND_API_KEY", "")
FROM_EMAIL              = os.environ.get("FROM_EMAIL", "info@mail.harborprivacy.com")
BASE_URL                = os.environ.get("BASE_URL", "https://fax.harborprivacy.com")

PRICE_BASE              = "price_1TTY2iCOrGNrBgIfVhhszJMO"
PRICE_EXTRA_PAGES       = "price_1TTY4CCOrGNrBgIfEr8byUPO"
PRICE_REMOVE_BRANDING   = "price_1TTYM1COrGNrBgIfzaOhdSjz"

AMOUNT_BASE             = 299
AMOUNT_EXTRA_PAGES      = 199
AMOUNT_REMOVE_BRANDING  = 99
PAGE_LIMIT_BASE         = 10
PAGE_LIMIT_EXTRA        = 30
MAX_FILE_MB             = 20

UPLOAD_DIR = Path("/tmp/harbor-fax-uploads")
MEDIA_DIR  = Path("/var/www/network/fax-media")
DB_PATH    = "/home/ubuntu/harbor-fax.db"

UPLOAD_DIR.mkdir(exist_ok=True)
MEDIA_DIR.mkdir(exist_ok=True)

stripe.api_key = STRIPE_SECRET
telnyx_client = telnyx.Telnyx(api_key=TELNYX_API_KEY)

ALLOWED_EXTS = {"pdf", "jpg", "jpeg", "png", "docx"}

_UPLOAD_RATE: dict = {}  # {ip: [timestamps]}

def _check_upload_rate(ip: str) -> bool:
    import time as _t
    now = _t.time()
    hits = [t for t in _UPLOAD_RATE.get(ip, []) if now - t < 60]
    _UPLOAD_RATE[ip] = hits
    if len(hits) >= 5:
        return False
    _UPLOAD_RATE[ip].append(now)
    return True


# ---------------------------------------------------------------------------
# Database
# ---------------------------------------------------------------------------

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_db() as db:
        db.executescript("""
            CREATE TABLE IF NOT EXISTS fax_files (
                token       TEXT PRIMARY KEY,
                orig_name   TEXT,
                file_path   TEXT,
                page_count  INTEGER,
                created_at  TEXT
            );
            CREATE TABLE IF NOT EXISTS fax_orders (
                token           TEXT PRIMARY KEY,
                status          TEXT,
                fax_number      TEXT,
                to_name         TEXT,
                from_name       TEXT,
                subject         TEXT,
                message         TEXT,
                remove_branding INTEGER,
                extra_pages     INTEGER,
                email           TEXT,
                file_tokens     TEXT,
                amount          INTEGER,
                promo_id        TEXT,
                payment_intent  TEXT,
                telnyx_fax_id   TEXT,
                merged_pdf      TEXT,
                created_at      TEXT,
                updated_at      TEXT
            );
        """)
        for col, ddl in (
            ("retry_count",   "ALTER TABLE fax_orders ADD COLUMN retry_count INTEGER DEFAULT 0"),
            ("refunded",      "ALTER TABLE fax_orders ADD COLUMN refunded INTEGER DEFAULT 0"),
            ("refund_id",     "ALTER TABLE fax_orders ADD COLUMN refund_id TEXT"),
            ("last_failure",  "ALTER TABLE fax_orders ADD COLUMN last_failure TEXT"),
        ):
            try: db.execute(ddl)
            except sqlite3.OperationalError: pass

init_db()

MAX_RETRIES   = 2     # total attempts = 1 initial + MAX_RETRIES
RETRY_DELAY_S = 60


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def ntfy(title, body, priority="default", tags="fax"):
    try:
        requests.post(
            "https://ntfy.harborprivacy.com/harbor-alerts",
            data=body.encode(),
            headers={"Title": title, "Priority": priority, "Tags": tags,
                     **({"Authorization": f"Basic {NTFY_AUTH}"} if NTFY_AUTH else {})},
            timeout=5,
        )
    except Exception:
        pass


def count_pdf_pages(path):
    try:
        reader = PyPDF2.PdfReader(str(path))
        return len(reader.pages)
    except Exception:
        return 1


def image_to_pdf_bytes(img_path):
    from reportlab.lib.utils import ImageReader
    img = Image.open(img_path)
    try:
        from PIL import ImageOps
        img = ImageOps.exif_transpose(img)
    except Exception:
        pass
    img = img.convert("RGB")

    page_w, page_h = letter  # always portrait letter
    margin = 0.25 * inch
    max_w = page_w - 2 * margin
    max_h = page_h - 2 * margin

    iw, ih = img.size
    scale = min(max_w / iw, max_h / ih)
    draw_w, draw_h = iw * scale, ih * scale
    x = (page_w - draw_w) / 2
    y = (page_h - draw_h) / 2

    buf = BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    c.drawImage(ImageReader(img), x, y, width=draw_w, height=draw_h,
                preserveAspectRatio=True, mask='auto')
    c.showPage()
    c.save()
    return buf.getvalue()


def docx_to_pdf_bytes(docx_path):
    doc = Document(str(docx_path))
    buf = BytesIO()
    story = []
    styles = getSampleStyleSheet()
    body_style = styles["BodyText"]
    body_style.fontSize = 11
    body_style.leading = 14

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 0.1 * inch))
            continue
        story.append(Paragraph(text, body_style))
        story.append(Spacer(1, 0.05 * inch))

    pdf = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    pdf.build(story)
    return buf.getvalue()


def estimate_docx_pages(docx_path):
    try:
        doc = Document(str(docx_path))
        total_lines = sum(
            max(1, len(p.text) // 80 + 1) for p in doc.paragraphs if p.text.strip()
        )
        return max(1, (total_lines + 49) // 50)
    except Exception:
        return 1


def make_cover_page_bytes(to_name="", fax_number="", from_name="", subject="", message="", page_count=0, branded=True):
    buf = BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    w, h = letter

    # White background
    c.setFillColorRGB(1, 1, 1)
    c.rect(0, 0, w, h, fill=1, stroke=0)

    INK         = (0.10, 0.14, 0.18)   # near-black for headings/body
    MUTED       = (0.35, 0.40, 0.45)   # mid grey for labels
    DIVIDER     = (0.78, 0.80, 0.82)   # light grey divider
    ACCENT      = (0.12, 0.36, 0.42)   # harbor teal for brand title

    if branded:
        c.setFillColorRGB(*ACCENT)
        c.setFont("Helvetica-Bold", 22)
        c.drawCentredString(w / 2, h - 1.2 * inch, "Harbor Privacy Fax")
        c.setFillColorRGB(*MUTED)
        c.setFont("Helvetica", 10)
        c.drawCentredString(w / 2, h - 1.6 * inch, "fax.harborprivacy.com")
        divider_y = h - 1.9 * inch
        fields_y = h - 2.4 * inch
    else:
        c.setFillColorRGB(*INK)
        c.setFont("Helvetica-Bold", 18)
        c.drawCentredString(w / 2, h - 1.2 * inch, "FAX COVER SHEET")
        divider_y = h - 1.5 * inch
        fields_y = h - 2.0 * inch

    # Divider
    c.setFillColorRGB(*DIVIDER)
    c.rect(inch, divider_y, w - 2 * inch, 0.015 * inch, fill=1, stroke=0)

    # Cover fields
    label_x = inch
    value_x = 2.5 * inch
    y = fields_y
    line_h = 0.35 * inch

    fields = [
        ("To:", to_name or "—"),
        ("Fax Number:", fax_number or "—"),
        ("From:", from_name or "—"),
        ("Date:", datetime.utcnow().strftime("%B %d, %Y")),
        ("Pages:", str(page_count + 1) + " (including cover)"),
        ("Subject:", subject or "—"),
    ]

    for label, value in fields:
        c.setFillColorRGB(*MUTED)
        c.setFont("Helvetica-Bold", 11)
        c.drawString(label_x, y, label)
        c.setFillColorRGB(*INK)
        c.setFont("Helvetica", 11)
        c.drawString(value_x, y, value[:70])
        y -= line_h

    # Message box
    if message:
        y -= 0.1 * inch
        c.setFillColorRGB(*DIVIDER)
        c.rect(inch, y - 0.015 * inch, w - 2 * inch, 0.015 * inch, fill=1, stroke=0)
        y -= 0.35 * inch
        c.setFillColorRGB(*MUTED)
        c.setFont("Helvetica-Bold", 11)
        c.drawString(label_x, y, "Message:")
        y -= 0.3 * inch
        c.setFillColorRGB(*INK)
        c.setFont("Helvetica", 10)
        for line in message.splitlines():
            for chunk in [line[i:i+90] for i in range(0, max(len(line), 1), 90)]:
                c.drawString(label_x, y, chunk)
                y -= 0.22 * inch
                if y < 1.2 * inch:
                    break

    if branded:
        c.setFillColorRGB(*DIVIDER)
        c.rect(inch, 0.9 * inch, w - 2 * inch, 0.015 * inch, fill=1, stroke=0)
        c.setFillColorRGB(*MUTED)
        c.setFont("Helvetica", 9)
        c.drawCentredString(w / 2, 0.65 * inch,
                            "Harbor Privacy | HIPAA Conduit Exception | Operated under U.S. law")
        c.drawCentredString(w / 2, 0.45 * inch,
                            "No sender identity is stored. Document deleted on delivery.")

    c.save()
    return buf.getvalue()


def file_to_pdf_bytes(file_path, ext):
    ext = ext.lower().lstrip(".")
    if ext == "pdf":
        return Path(file_path).read_bytes()
    if ext in ("jpg", "jpeg", "png"):
        return image_to_pdf_bytes(file_path)
    if ext == "docx":
        return docx_to_pdf_bytes(file_path)
    raise ValueError(f"Unsupported extension: {ext}")


def merge_pdfs(pdf_bytes_list):
    writer = PyPDF2.PdfWriter()
    for pdf_bytes in pdf_bytes_list:
        reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)
    buf = BytesIO()
    writer.write(buf)
    return buf.getvalue()


def build_fax_pdf(order_token, file_tokens_list, remove_branding, order=None):
    db = get_db()
    pdf_parts = []

    total_pages = 0
    for ft in file_tokens_list:
        row = db.execute("SELECT page_count FROM fax_files WHERE token=?", (ft,)).fetchone()
        if row:
            total_pages += row["page_count"] or 1
    cover_kwargs = {}
    if order:
        cover_kwargs = dict(
            to_name=order["to_name"] or "",
            fax_number=order["fax_number"] or "",
            from_name=order["from_name"] or "",
            subject=order["subject"] or "",
            message=order["message"] or "",
            page_count=total_pages,
        )
    pdf_parts.append(make_cover_page_bytes(**cover_kwargs, branded=not remove_branding))

    for ft in file_tokens_list:
        row = db.execute("SELECT file_path, orig_name FROM fax_files WHERE token=?", (ft,)).fetchone()
        if not row:
            continue
        fpath = row["file_path"]
        ext = Path(row["orig_name"]).suffix.lstrip(".")
        pdf_parts.append(file_to_pdf_bytes(fpath, ext))

    db.close()
    merged = merge_pdfs(pdf_parts)
    out_path = MEDIA_DIR / f"{order_token}.pdf"
    out_path.write_bytes(merged)
    return str(out_path)


def send_fax_async(order_token):
    threading.Thread(target=_send_fax, args=(order_token,), daemon=True).start()


def _send_fax(order_token):
    db = get_db()
    try:
        order = db.execute("SELECT * FROM fax_orders WHERE token=?", (order_token,)).fetchone()
        if not order:
            return

        file_tokens = json.loads(order["file_tokens"])
        merged_path = build_fax_pdf(order_token, file_tokens, bool(order["remove_branding"]), order=order)
        media_url = f"{BASE_URL}/fax-media/{order_token}.pdf"

        fax_number = order["fax_number"]
        if not fax_number.startswith("+"):
            fax_number = "+1" + fax_number.replace("-", "").replace(" ", "").replace("(", "").replace(")", "")

        resp = telnyx_client.faxes.create(
            connection_id=TELNYX_CONNECTION_ID,
            from_=TELNYX_FROM_NUMBER,
            to=fax_number,
            media_url=media_url,
            store_media=False,
        )

        telnyx_fax_id = resp.data.id
        db.execute(
            "UPDATE fax_orders SET status='sending', telnyx_fax_id=?, merged_pdf=?, updated_at=? WHERE token=?",
            (telnyx_fax_id, merged_path, datetime.utcnow().isoformat(), order_token),
        )
        db.commit()
        log.info("Fax %s queued via Telnyx as %s", order_token, telnyx_fax_id)

    except Exception as e:
        log.exception("Fax send failed for %s", order_token)
        db.execute(
            "UPDATE fax_orders SET status='failed', updated_at=? WHERE token=?",
            (datetime.utcnow().isoformat(), order_token),
        )
        db.commit()
        log.error("Fax send error for %s: %s", order_token, e)
    finally:
        db.close()


def delete_fax_files(order_token):
    db = get_db()
    try:
        order = db.execute("SELECT file_tokens, merged_pdf FROM fax_orders WHERE token=?", (order_token,)).fetchone()
        if not order:
            return
        for ft in json.loads(order["file_tokens"] or "[]"):
            row = db.execute("SELECT file_path FROM fax_files WHERE token=?", (ft,)).fetchone()
            if row:
                try:
                    Path(row["file_path"]).unlink(missing_ok=True)
                except Exception:
                    pass
        if order["merged_pdf"]:
            try:
                Path(order["merged_pdf"]).unlink(missing_ok=True)
            except Exception:
                pass
    finally:
        db.close()


def _resend_fax(order_token):
    """Re-send an existing order to Telnyx using the already-merged PDF."""
    db = get_db()
    try:
        order = db.execute("SELECT * FROM fax_orders WHERE token=?", (order_token,)).fetchone()
        if not order:
            return
        merged_path = order["merged_pdf"]
        if not merged_path or not Path(merged_path).exists():
            file_tokens = json.loads(order["file_tokens"] or "[]")
            merged_path = build_fax_pdf(order_token, file_tokens, bool(order["remove_branding"]), order=order)
        media_url = f"{BASE_URL}/fax-media/{order_token}.pdf"
        fax_number = order["fax_number"]
        if not fax_number.startswith("+"):
            fax_number = "+1" + "".join(c for c in fax_number if c.isdigit())[-10:]
        resp = telnyx_client.faxes.create(
            connection_id=TELNYX_CONNECTION_ID,
            from_=TELNYX_FROM_NUMBER,
            to=fax_number,
            media_url=media_url,
            store_media=False,
        )
        new_id = resp.data.id
        db.execute(
            "UPDATE fax_orders SET status='sending', telnyx_fax_id=?, merged_pdf=?, updated_at=? WHERE token=?",
            (new_id, merged_path, datetime.utcnow().isoformat(), order_token),
        )
        db.commit()
        log.info("Fax %s retry queued as %s (attempt %s)", order_token, new_id, (order["retry_count"] or 0) + 1)
    except Exception:
        log.exception("Retry send failed for %s", order_token)
        db.execute(
            "UPDATE fax_orders SET status='failed', updated_at=? WHERE token=?",
            (datetime.utcnow().isoformat(), order_token),
        )
        db.commit()
    finally:
        db.close()


def schedule_retry(order_token, delay_s=RETRY_DELAY_S):
    """Bump retry_count and schedule a re-send after delay_s seconds."""
    db = get_db()
    try:
        db.execute(
            "UPDATE fax_orders SET retry_count = COALESCE(retry_count,0) + 1, status='retrying', updated_at=? WHERE token=?",
            (datetime.utcnow().isoformat(), order_token),
        )
        db.commit()
    finally:
        db.close()
    t = threading.Timer(delay_s, _resend_fax, args=(order_token,))
    t.daemon = True
    t.start()


def refund_order(order_token, reason="fax delivery failed"):
    """Refund the Stripe PaymentIntent (idempotent) and mark order refunded."""
    db = get_db()
    try:
        order = db.execute(
            "SELECT token, payment_intent, refunded, amount, email, fax_number FROM fax_orders WHERE token=?",
            (order_token,),
        ).fetchone()
        if not order:
            return False
        if order["refunded"]:
            return True
        pi = order["payment_intent"]
        if not pi:
            log.warning("Cannot refund %s: no payment_intent on order", order_token)
            return False
        try:
            refund = stripe.Refund.create(
                payment_intent=pi,
                reason="requested_by_customer",
                metadata={"order_token": order_token, "reason": reason},
            )
            db.execute(
                "UPDATE fax_orders SET refunded=1, refund_id=?, updated_at=? WHERE token=?",
                (refund.id, datetime.utcnow().isoformat(), order_token),
            )
            db.commit()
            log.info("Refunded %s for order %s (refund %s)", order["amount"], order_token, refund.id)
            return True
        except stripe.error.StripeError:
            log.exception("Stripe refund failed for %s", order_token)
            return False
    finally:
        db.close()


def send_delivery_email(email, order_token, fax_number, status, refunded=False):
    if not email or not RESEND_API_KEY:
        return
    if status == "delivered":
        subject = "Your fax was delivered"
        body = (
            f"Your fax to {fax_number} has been delivered successfully.\n\n"
            f"Order reference: {order_token[:8].upper()}\n\n"
            "As promised, your document has been permanently deleted.\n\n"
            "Harbor Privacy Fax | fax.harborprivacy.com"
        )
    else:
        subject = "Your fax could not be delivered"
        refund_line = (
            "Your payment has been refunded in full and should appear on your card within 5-10 business days.\n\n"
            if refunded else
            "If your card was charged, your payment will be refunded automatically.\n\n"
        )
        body = (
            f"Your fax to {fax_number} could not be delivered after multiple attempts.\n\n"
            f"{refund_line}"
            f"Order reference: {order_token[:8].upper()}\n\n"
            "As promised, your document has been permanently deleted.\n\n"
            "Harbor Privacy Fax | fax.harborprivacy.com"
        )
    try:
        requests.post(
            "https://api.resend.com/emails",
            headers={"Authorization": f"Bearer {RESEND_API_KEY}", "Content-Type": "application/json"},
            json={"from": FROM_EMAIL, "to": email, "subject": subject, "text": body},
            timeout=10,
        )
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/fax/upload", methods=["POST"])
def fax_upload():
    ip = request.headers.get("X-Forwarded-For", request.remote_addr).split(",")[0].strip()
    if not _check_upload_rate(ip):
        return jsonify({"error": "Too many uploads. Please wait a moment."}), 429
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file provided"}), 400

    ext = Path(f.filename or "").suffix.lstrip(".").lower()
    if not ext or ext not in ALLOWED_EXTS:
        return jsonify({"error": "Unsupported file type. Use PDF, JPG, PNG, or DOCX."}), 400

    data = f.read()
    if len(data) > MAX_FILE_MB * 1024 * 1024:
        return jsonify({"error": f"File exceeds {MAX_FILE_MB}MB limit."}), 400

    token = uuid.uuid4().hex
    fpath = UPLOAD_DIR / f"{token}.{ext}"
    fpath.write_bytes(data)

    if ext == "pdf":
        pages = count_pdf_pages(fpath)
    elif ext in ("jpg", "jpeg", "png"):
        pages = 1
    elif ext == "docx":
        pages = estimate_docx_pages(fpath)
    else:
        pages = 1

    db = get_db()
    db.execute(
        "INSERT INTO fax_files (token, orig_name, file_path, page_count, created_at) VALUES (?,?,?,?,?)",
        (token, f.filename, str(fpath), pages, datetime.utcnow().isoformat()),
    )
    db.commit()
    db.close()
    return jsonify({"file_token": token, "page_count": pages})


def _get_coupon_discount(promo_id_or_code, amount, is_code=False):
    if is_code:
        promos = stripe.PromotionCode.list(code=promo_id_or_code, active=True, limit=1)
        if not promos.data:
            return None, None, "Invalid or expired promo code"
        promo = promos.data[0]
        promo_id = promo.id
    else:
        promo = stripe.PromotionCode.retrieve(promo_id_or_code)
        promo_id = promo_id_or_code
    d = promo.to_dict()
    coupon_id = (d.get("promotion") or {}).get("coupon") or d.get("coupon")
    if not coupon_id:
        return None, None, "Promo has no coupon"
    coupon = stripe.Coupon.retrieve(coupon_id) if isinstance(coupon_id, str) else coupon_id
    cd = coupon.to_dict() if hasattr(coupon, "to_dict") else coupon
    if cd.get("percent_off"):
        discount = int(round(amount * cd["percent_off"] / 100))
    elif cd.get("amount_off"):
        discount = cd["amount_off"]
    else:
        return None, None, "Promo has no valid discount"
    return promo_id, min(discount, amount), None


@app.route("/fax/validate-promo", methods=["POST"])
def validate_promo():
    body = request.get_json(silent=True) or {}
    code = (body.get("code") or "").strip().upper()
    if not code:
        return jsonify({"error": "No code provided"}), 400
    try:
        promos = stripe.PromotionCode.list(code=code, active=True, limit=1)
        if not promos.data:
            return jsonify({"error": "Invalid or expired promo code"}), 400
        amount = body.get("amount", 299)
        promo_id, discount, err = _get_coupon_discount(code, amount, is_code=True)
        if err:
            return jsonify({"error": err}), 400
        return jsonify({"promo_id": promo_id, "discount": discount})
    except stripe.error.StripeError as e:
        return jsonify({"error": str(e)}), 400


def _calc_amount(remove_branding, extra_pages):
    return AMOUNT_BASE + (AMOUNT_REMOVE_BRANDING if remove_branding else 0) + (AMOUNT_EXTRA_PAGES if extra_pages else 0)


def _validate_order_payload(body):
    fax_number = (body.get("fax_number") or "").strip()
    if not fax_number:
        return None, "Fax number is required"
    digits = "".join(c for c in fax_number if c.isdigit())
    if len(digits) == 11 and digits.startswith("1"):
        digits = digits[1:]
    if len(digits) != 10:
        return None, "Fax number must be a 10-digit US number"
    body["fax_number"] = "+1" + digits
    file_tokens = body.get("file_tokens") or []
    if not file_tokens:
        return None, "No files attached"
    db = get_db()
    for ft in file_tokens:
        row = db.execute("SELECT token FROM fax_files WHERE token=?", (ft,)).fetchone()
        if not row:
            db.close()
            return None, f"File token not found: {ft}"
    db.close()
    return True, None


@app.route("/fax/create-payment-intent", methods=["POST"])
def create_payment_intent():
    body = request.get_json(silent=True) or {}
    ok, err = _validate_order_payload(body)
    if not ok:
        return jsonify({"error": err}), 400

    remove_branding = bool(body.get("remove_branding"))
    extra_pages     = bool(body.get("extra_pages"))
    promo_id        = body.get("promo_id")
    discount        = int(body.get("discount") or 0)
    email           = (body.get("email") or "").strip()

    amount = _calc_amount(remove_branding, extra_pages)
    if promo_id and discount:
        try:
            _, real_discount, err = _get_coupon_discount(promo_id, amount)
            if err:
                return jsonify({"error": err}), 400
            amount = max(0, amount - real_discount)
        except Exception:
            return jsonify({"error": "Invalid promo code"}), 400

    if amount == 0:
        return jsonify({"error": "Use /fax/send-free for free orders"}), 400

    order_token = body.get("file_token") or uuid.uuid4().hex

    meta = {
        "order_token": order_token,
        "fax_number":  body.get("fax_number", ""),
        "service":     "harbor-fax",
    }

    try:
        intent = stripe.PaymentIntent.create(
            amount=amount,
            currency="usd",
            receipt_email=email or None,
            metadata=meta,
        )
    except stripe.error.StripeError as e:
        return jsonify({"error": str(e)}), 400

    db = get_db()
    db.execute(
        """INSERT OR REPLACE INTO fax_orders
           (token, status, fax_number, to_name, from_name, subject, message,
            remove_branding, extra_pages, email, file_tokens, amount, promo_id,
            payment_intent, created_at, updated_at)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            order_token, "pending_payment",
            body.get("fax_number"), body.get("to_name"), body.get("from_name"),
            body.get("subject"), body.get("message"),
            int(remove_branding), int(extra_pages), email,
            json.dumps(body.get("file_tokens", [])),
            amount, promo_id, intent.id,
            datetime.utcnow().isoformat(), datetime.utcnow().isoformat(),
        ),
    )
    db.commit()
    db.close()

    return jsonify({"client_secret": intent.client_secret})


@app.route("/fax/send-free", methods=["POST"])
def send_free():
    body = request.get_json(silent=True) or {}
    ok, err = _validate_order_payload(body)
    if not ok:
        return jsonify({"error": err}), 400

    promo_id = body.get("promo_id")
    discount = int(body.get("discount") or 0)
    remove_branding = bool(body.get("remove_branding"))
    extra_pages     = bool(body.get("extra_pages"))
    amount = _calc_amount(remove_branding, extra_pages)

    if promo_id and discount:
        try:
            _, real_discount, err = _get_coupon_discount(promo_id, amount)
            if err:
                return jsonify({"error": err}), 400
            if real_discount < amount:
                return jsonify({"error": "Promo does not cover full amount"}), 400
        except Exception:
            return jsonify({"error": "Invalid promo code"}), 400
    else:
        return jsonify({"error": "No valid promo code for free send"}), 400

    order_token = body.get("file_token") or uuid.uuid4().hex
    email = (body.get("email") or "").strip()

    db = get_db()
    db.execute(
        """INSERT OR REPLACE INTO fax_orders
           (token, status, fax_number, to_name, from_name, subject, message,
            remove_branding, extra_pages, email, file_tokens, amount, promo_id,
            created_at, updated_at)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            order_token, "queued",
            body.get("fax_number"), body.get("to_name"), body.get("from_name"),
            body.get("subject"), body.get("message"),
            int(remove_branding), int(extra_pages), email,
            json.dumps(body.get("file_tokens", [])),
            0, promo_id,
            datetime.utcnow().isoformat(), datetime.utcnow().isoformat(),
        ),
    )
    db.commit()
    db.close()

    send_fax_async(order_token)
    return jsonify({"success": True})


@app.route("/fax/status/<token>")
def fax_status(token):
    db = get_db()
    order = db.execute(
        "SELECT status, fax_number, to_name, created_at, updated_at, telnyx_fax_id FROM fax_orders WHERE token=?",
        (token,),
    ).fetchone()
    db.close()
    if not order:
        return jsonify({"error": "Order not found"}), 404
    return jsonify(dict(order))


@app.route("/fax/status-page")
def fax_status_page():
    token = request.args.get("token", "")
    return redirect(f"/fax-status.html?token={token}")


@app.route("/fax/stripe-webhook", methods=["POST"])
def stripe_webhook():
    payload = request.get_data()
    sig = request.headers.get("Stripe-Signature", "")
    try:
        if STRIPE_WEBHOOK_SECRET:
            event = stripe.Webhook.construct_event(payload, sig, STRIPE_WEBHOOK_SECRET)
        else:
            event = stripe.Event.construct_from(json.loads(payload), stripe.api_key)
    except Exception as e:
        log.warning("Stripe webhook error: %s", e)
        return "Bad request", 400

    if event["type"] == "payment_intent.succeeded":
        pi = event["data"]["object"]
        order_token = pi.get("metadata", {}).get("order_token")
        if order_token:
            db = get_db()
            db.execute(
                "UPDATE fax_orders SET status='queued', updated_at=? WHERE token=? AND status='pending_payment'",
                (datetime.utcnow().isoformat(), order_token),
            )
            db.commit()
            db.close()
            send_fax_async(order_token)

    return "ok", 200


@app.route("/fax/telnyx-webhook", methods=["POST"])
def telnyx_webhook():
    try:
        body = request.get_json(force=True) or {}
    except Exception:
        return "ok", 200

    event_type = body.get("data", {}).get("event_type", "")
    payload    = body.get("data", {}).get("payload", {})
    fax_id     = payload.get("fax_id") or payload.get("id", "")
    status     = payload.get("status", "")

    log.info("Telnyx event: %s fax_id=%s status=%s", event_type, fax_id, status)

    if not fax_id:
        return "ok", 200

    db = get_db()
    order = db.execute("SELECT * FROM fax_orders WHERE telnyx_fax_id=?", (fax_id,)).fetchone()
    if not order:
        db.close()
        return "ok", 200

    if status == "delivered":
        db.execute(
            "UPDATE fax_orders SET status='delivered', updated_at=? WHERE telnyx_fax_id=?",
            (datetime.utcnow().isoformat(), fax_id),
        )
        db.commit()
        db.close()
        delete_fax_files(order["token"])
        send_delivery_email(order["email"], order["token"], order["fax_number"], "delivered")
        return "ok", 200

    if status == "failed":
        failure_reason = payload.get("failure_reason") or payload.get("failure_details") or ""
        retries = order["retry_count"] or 0
        db.execute(
            "UPDATE fax_orders SET last_failure=?, updated_at=? WHERE telnyx_fax_id=?",
            (str(failure_reason)[:500], datetime.utcnow().isoformat(), fax_id),
        )
        db.commit()
        order_token = order["token"]
        db.close()

        if retries < MAX_RETRIES:
            log.info("Fax %s failed (attempt %s/%s), scheduling retry in %ss. Reason: %s",
                     order_token, retries + 1, MAX_RETRIES + 1, RETRY_DELAY_S, failure_reason)
            schedule_retry(order_token, RETRY_DELAY_S)
            return "ok", 200

        log.warning("Fax %s permanently failed after %s attempts. Reason: %s",
                    order_token, retries + 1, failure_reason)
        refunded = refund_order(order_token, reason=str(failure_reason)[:200])
        delete_fax_files(order_token)
        send_delivery_email(order["email"], order_token, order["fax_number"], "failed", refunded=refunded)
        return "ok", 200

    db.close()
    return "ok", 200


@app.route("/health")
@app.route("/fax/health")
def fax_health():
    try:
        stripe.Balance.retrieve()
        stripe_ok = True
    except Exception:
        stripe_ok = False
    status = "ok" if stripe_ok else "degraded"
    code = 200 if stripe_ok else 503
    return jsonify({"status": status, "stripe": "ok" if stripe_ok else "error"}), code


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=7500)
